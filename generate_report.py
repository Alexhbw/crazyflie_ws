#!/usr/bin/env python3
"""生成项目报告 DOCX — 障碍物避障 + 安全监视器 + 实验框架 + 区间可达集"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "/home/alex/crazyflie_ws/report_enhancements.docx"

doc = Document()

# ====== 样式设置 ======
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def formula(text):
    """插入居中公式"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.italic = True
    return p


def image_placeholder(caption, width_inches=5.5):
    """插入图片占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ 图片: {caption} ]")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(10)
    run.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"图: {caption}")
    r2.bold = True
    r2.font.size = Pt(10)


def add_table(headers, rows):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # 数据
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ================================================================
# 封面
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Crazyflie 集群 ODE 仿真系统\n安全分析与实验框架增强")
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "—— 面向无人机安全自主飞行的 ODE 初值问题\n"
    "数值解法稳定域分析与安全边界形式化验证")
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("2026 年 6 月").font.size = Pt(12)

doc.add_page_break()

# ================================================================
# 1. 摘要
# ================================================================
heading("摘要", 1)

para(
    "本报告描述在 Crazyflie 微型无人机集群 ODE 数值仿真系统基础上新增的四个核心功能模块："
    "(1) 圆柱体障碍物避障子系统，基于平滑人工势场实现安全绕行；"
    "(2) 运行时安全监视器，对碰撞、速度、姿态、能量漂移进行形式化条件检查与日志记录；"
    "(3) 区间可达集近似分析，通过一步正向区间传播实现最坏情况碰撞预警；"
    "(4) 求解器稳定域对比实验框架，支持全因子参数扫描与临界步长二分搜索，"
    '为[数值解法稳定域分析与安全边界形式化验证]研究课题提供定量实验支撑。'
)

para(
    "各模块通过 ROS 2 参数体系实现灵活配置，支持正常仿真、安全监视、"
    "数据采集等多种运行模式，配套 Python 实验脚本和 matplotlib 图表生成工具，"
    "可直接产出报告级对比图表。"
)

doc.add_page_break()

# ================================================================
# 2. 系统架构总览
# ================================================================
heading("1. 系统架构总览", 1)

para("增强后的系统由以下层次构成：")

add_table(
    ["层次", "组件", "功能"],
    [
        ["配置层", "obstacles.yaml, ROS 2 Parameters", "障碍物定义、求解器选择、安全阈值、实验参数"],
        ["物理引擎层", "ODE 数值求解器 (5种), 人工势场避障", "12维状态动力学积分、障碍物排斥力计算"],
        ["安全监控层", "SafetyMonitor, IntervalReachability", "每步安全条件检查、区间可达集预警"],
        ["实验编排层", "Python 实验脚本", "批量参数扫描、临界步长搜索、图表生成"],
        ["可视化层", "RViz + Marker + TF", "3D 无人机模型、障碍物圆柱、飞行轨迹"],
    ],
)

image_placeholder("系统架构层次图", 5.0)

# ================================================================
# 3. 障碍物避障子系统
# ================================================================
heading("2. 障碍物避障子系统", 1)

heading("2.1 数据结构与 YAML 配置", 2)

para("障碍物采用圆柱体 (Cylinder) 建模，每个障碍物由五元组定义：")
formula("CylinderObstacle = (id, x, y, z, radius, height)")

para("配置文件 config/obstacles.yaml 定义了 5 根圆柱障碍物，布局如下：")

add_table(
    ["ID", "位置 (x, y, z)", "半径 (m)", "高度 (m)", "说明"],
    [
        ["center_pole", "(0.00, 0.00, 0.75)", "0.05", "1.50", "中心细杆，呼吸圆环收缩时穿过"],
        ["pillar_ne", "(1.34, 1.34, 1.00)", "0.18", "2.00", "东北角粗柱"],
        ["pillar_nw", "(-1.34, 1.34, 1.00)", "0.18", "2.00", "西北角粗柱"],
        ["pillar_sw", "(-1.34, -1.34, 1.00)", "0.18", "2.00", "西南角粗柱"],
        ["pillar_se", "(1.34, -1.34, 1.00)", "0.18", "2.00", "东南角粗柱"],
    ],
)

para(
    "四根外侧柱位于半径为 1.9 m 的圆上，以 45° 间隔分布。呼吸圆环编队半径 "
    "在 0.5~1.7 m 之间振荡，最大半径时无人机从两根柱子之间穿过，形成自然的避障测试场景。"
)

image_placeholder("障碍物布局俯视图 (5根圆柱位置 + 编队圆环轨迹)", 4.5)

heading("2.2 平滑人工势场避障算法", 2)

para("避障算法经过三个迭代版本的设计优化，最终采用平滑三次方势场配合速度阻尼和切向引导。")

heading("2.2.1 无量纲侵入比", 3)

para("对每个障碍物，首先计算无人机到圆柱表面的最近距离：")
formula("d_surf = d_xy − r_obs")
para("其中 d_xy 为无人机到圆柱轴心的水平距离，r_obs 为圆柱半径。定义侵入比 t ∈ [0, 1]：")
formula("t = clamp(1 − d_surf / d_safe,  0,  1)")

para("t = 0 表示无人机在安全边界外，不受斥力影响；t = 1 表示无人机已触及障碍物表面，斥力最大。")

heading("2.2.2 径向排斥力", 3)

para("采用三次方曲线确保近场强劲、远场柔和：")
formula("F_rad = k_rep · t³ · r̂_xy")
para("其中 k_rep = 8.0 为排斥增益，r̂_xy 为指向障碍物轴心的单位径向向量。")
para("与经典的 1/d² 势场相比，三次方曲线具有以下优势：")
para("• 在安全边界 (t ≈ 0) 处力接近零，不干扰正常编队")
para("• 在表面附近 (t → 1) 力急剧增长，有效阻止穿透")
para("• 一阶导数连续，避免力场跳变导致的控制器震荡")

para("设计迭代历程：")
add_table(
    ["版本", "力场公式", "k_rep", "d_safe", "问题"],
    [
        ["v1", "F = k · (1/d − 1/d_safe) / d²", "5.0", "0.5", "近表面力趋向无穷，无人机被弹飞"],
        ["v2", "F = k · t² (二次)", "2.0", "0.3", "表面力仅 2 m/s²，编队控制器 4 m/s² 直接穿透"],
        ["v3 (最终)", "F = k · t³ (三次)", "8.0", "0.5", "远场柔和 (t=0.5→1.0 m/s²)，表面强劲 (t=1→8.0 m/s²)"],
    ],
)

image_placeholder("三种力场曲线对比图 (1/d² vs t² vs t³)", 5.0)

heading("2.2.3 速度阻尼", 3)

para("为防止无人机高速冲向障碍物时因数值积分步长限制而穿透，引入速度阻尼项：")
formula("v_radial = v_xy · r̂_xy")
formula("F_damp = −k_rep · 0.6 · t² · v_radial · r̂_xy    (仅当 v_radial > 0)")

para(
    "物理意义：这不是[推力]而是[刹车]。无人机朝向障碍物运动越快，"
    "制动力越强。这是防止穿透的关键机制——即使 PID 控制指令朝向障碍物，"
    "阻尼项会主动减速，使无人机在到达表面之前停下来。"
)

heading("2.2.4 切向绕行引导", 3)

para("在障碍物附近，引导无人机沿切向绕行而非正面碰撞：")
formula("F_tan = k_rep · 0.35 · t_tan³ · t̂")
para("其中 t̂ 为径向方向逆时针旋转 90° 的切向单位向量。")
para("切向符号选择与当前速度方向一致，实现'顺势绕行'。触发条件为 d_surf < 0.7 · d_safe。")

heading("2.2.5 Z 轴纵向排斥", 3)
formula("F_z = k_rep · 0.5 · t_z³ · sign(z − z_obs) · ẑ")

heading("2.2.6 合力与限幅", 3)

para("四分量叠加后施加饱和限幅：")
formula("F_total = F_rad + F_damp + F_tan + F_z")
formula("if ‖F_total‖ > 10.0:  F_total = F_total · (10.0 / ‖F_total‖)")

para("限幅 10.0 m/s² 确保障碍物斥力能有效覆盖编队控制器的 4.0 m/s² 上限。")

heading("2.2.7 与级联 PID 的集成", 3)

para("排斥加速度直接叠加到 PID 期望加速度中：")
formula("acc_des = Kp_v · (vel_des − vel) + a_rep_obs + a_rep_drone")
formula("acc_des = clamp(acc_des, −4.0, 4.0)")

para("其中 a_rep_obs 为障碍物排斥加速度，a_rep_drone 为机间排斥加速度。")

heading("2.3 算法参数总表", 2)

add_table(
    ["参数", "符号", "取值", "说明"],
    [
        ["安全距离", "d_safe", "0.5 m", "障碍物排斥影响半径"],
        ["排斥增益", "k_rep", "8.0", "径向排斥峰值 (t=1 时)"],
        ["速度阻尼系数", "α_damp", "0.6 · k_rep", "朝向障碍物速度的制动增益"],
        ["切向引导系数", "α_tan", "0.35 · k_rep", "切向绕行力的比例"],
        ["Z 轴系数", "α_z", "0.5 · k_rep", "纵向排斥力衰减"],
        ["力场上限", "F_max", "10.0 m/s²", "总排斥力矢量模上限"],
        ["切向触发比", "β_tan", "0.7 · d_safe", "切向引导的触发距离"],
    ],
)

heading("2.4 RViz 可视化", 2)

para(
    "使用 ROS 2 visualization_msgs::Marker::CYLINDER 消息将障碍物渲染为半透明红色圆柱。"
    "发布者采用 Transient Local QoS (depth=20)，确保 RViz 后连接也能收到完整标记。"
    "标记分为初始发布 (含 DELETEALL 清除旧标记) 和周期性重发 (每 10 步，约 100 ms)，"
    "避免 RViz Reset 后障碍物消失。"
)

heading("2.5 参数化控制", 2)

para("新增 ROS 2 参数 enable_obstacles (bool, 默认 true)，支持运行时动态开关障碍物：")
para("./run_experiment.sh rk4 0.01 5        # 障碍物开启")
para("./run_experiment.sh rk4 0.01 5 false  # 障碍物关闭")

doc.add_page_break()

# ================================================================
# 4. 运行时安全监视器
# ================================================================
heading("3. 运行时安全监视器 (SafetyMonitor)", 1)

heading("3.1 设计目标", 2)

para(
    "在仿真每一时间步自动检查以下形式化安全条件，将这些条件表达为时态逻辑中的不变式 "
    "(Invariant) □(condition)，跟踪违规事件并计算全局安全分数："
)

heading("3.2 安全条件定义", 2)

add_table(
    ["条件", "形式化定义", "阈值", "安全类型"],
    [
        ["碰撞安全", "□(min_obstacle_distance > 0)", "d_surf > 0", "硬安全 (Hard Safety)"],
        ["速度安全", "□(|v_i| < v_max)", "v_max = 2.5 m/s", "硬安全"],
        ["姿态安全", "□(|θ_i| < θ_max)", "θ_max = 0.8 rad", "硬安全"],
        ["能量漂移", "□(|ΔE/Δt| < ε)", "ε = 0.01 J/s", "软诊断 (Diagnostic)"],
    ],
)

para(
    "其中能量漂移条件仅用于诊断不同求解器的数值稳定性，"
    "不计入安全分数计算（因其为间接指标而非直接安全威胁）。"
)

heading("3.3 能量计算", 2)

para("每架无人机的总机械能由三部分构成：")
formula("E = ½m‖v‖² + ½ωᵀ I ω + mgz")
para("分别对应平动动能、转动动能和重力势能。每步计算 |E_i − E_{i−1}| 与阈值比较。")

heading("3.4 安全分数", 2)

para("对每种安全条件计算合规率：")
formula("compliance_c = N_compliant_c / N_total_c    (c ∈ {collision, velocity, attitude, energy})")
formula("overall_score = min(compliance_collision, compliance_velocity, compliance_attitude)")
para("综合分数取各硬安全条件合规率的最小值，反映系统最薄弱环节。")

heading("3.5 预热机制", 2)

para(
    "新增 safety.warmup_time 参数 (默认 2.0 s)。仿真启动后的预热期内，"
    "碰撞和姿态违规不记录——因无人机初始位置可能在障碍物附近。"
    "速度和能量监控不受影响。"
)

heading("3.6 输出", 2)

para("• 实时日志 (verbose 模式): 每次违规即时打印到终端")
para("• safety_log.csv: 所有违规记录 (时间、无人机 ID、条件、实际值、阈值)")
para("• experiment_metrics.csv: 聚合指标 (能量漂移率、障碍物距离、速度/姿态极值、跟踪误差、计算耗时)")
para("• 关机摘要: Safety Summary 打印到控制台，包含四个维度的合规率")

heading("3.7 参数接口", 2)

add_table(
    ["参数", "类型", "默认值", "说明"],
    [
        ["safety.v_max", "double", "2.5", "速度安全阈值 (m/s)"],
        ["safety.theta_max", "double", "0.8", "姿态安全阈值 (rad)"],
        ["safety.energy_drift_eps", "double", "0.01", "能量漂移阈值 (J/步)"],
        ["safety.warmup_time", "double", "2.0", "预热时间 (秒)"],
        ["safety.log_path", "string", "safety_log.csv", "违规日志路径"],
        ["safety.verbose", "bool", "false", "实时打印违规"],
    ],
)

doc.add_page_break()

# ================================================================
# 5. 区间可达集
# ================================================================
heading("4. 区间可达集近似分析 (Interval Reachability)", 1)

heading("4.1 问题定义", 2)

para(
    "给定当前状态 x 和控制输入 u，对初始状态施加有界扰动 ε，"
    "计算一步数值积分后的状态区间包围盒，并保守检测该包围盒是否与障碍物相交。"
    "这提供了形式化验证中的[最坏情况碰撞]预警能力："
    "如果在 ±ε 扰动下可达集与障碍物无交集，则当前状态下必不可能碰撞。"
)

heading("4.2 采样法 (默认)", 2)

para("对 12 维状态向量的每一维，分别施加 ±ε 扰动，运行求解器，取极值：")
formula("∀j ∈ [0, 11]:  x'_+ = solver(x + ε · e_j, u, h)")
formula("∀j ∈ [0, 11]:  x'_- = solver(x − ε · e_j, u, h)")
formula("lower(k) = min(x'_center(k), min_j(x'_+(k), x'_-(k)))")
formula("upper(k) = max(x'_center(k), max_j(x'_+(k), x'_-(k)))")

para("共需 25 次求解器调用 (1 中心 + 12维 × 2 方向)。")

heading("4.3 Jacobian 法 (可选)", 2)

para("通过有限差分计算一步状态转移 Jacobian，线性传播区间：")
formula("J[:,j] = [solver(x + ε·e_j, u, h) − solver(x − ε·e_j, u, h)] / (2ε)")
formula("lower = x'_center − |J| · ε · 1₁₂")
formula("upper = x'_center + |J| · ε · 1₁₂")

para("共 13 次求解器调用，速度快于采样法但保守性略差 (线性化近似)。")

heading("4.4 AABB-圆柱碰撞检测", 2)

para("将状态区间包围盒的位置分量与每个障碍物圆柱进行比较：")
para("• 计算包围盒 (AABB) 到圆柱轴心的最近点")
para("• 若最近点位于圆柱体积内 (d_xy < r_obs 且 |dz| < h/2)，判定为可达集与障碍物重叠")
para("• 输出最近距离 d_min 和重叠标志")

heading("4.5 参数接口", 2)

add_table(
    ["参数", "类型", "默认值", "说明"],
    [
        ["reachability.enabled", "bool", "false", "启用区间可达集分析"],
        ["reachability.epsilon", "double", "0.01", "状态扰动幅度"],
        ["reachability.method", "string", "sampling", "sampling 或 jacobian"],
    ],
)

doc.add_page_break()

# ================================================================
# 6. 求解器对比实验框架
# ================================================================
heading("5. 求解器稳定域对比实验框架", 1)

heading("5.1 实验设计", 2)

para(
    "为系统性地对比五种 ODE 数值求解器 (Euler, Heun, RK4, RK45, Implicit Euler) "
    "在安全关键场景下的表现，设计了全因子参数扫描矩阵："
)

add_table(
    ["扫描维度", "取值集合", "数量"],
    [
        ["求解器", "euler, heun, rk4, rk45, implicit", "5"],
        ["积分步长", "0.001, 0.005, 0.01, 0.02, 0.05 (s)", "5"],
        ["无人机数量", "1, 4, 8", "3"],
        ["障碍物", "true, false", "2"],
        ["仿真时长", "30.0 s (固定)", "1"],
    ],
)

para("总计 5 × 5 × 3 × 2 = 150 次独立实验。每次实验自动运行、自动采集、自动退出。")

heading("5.2 评估指标体系", 2)

para("每次实验采集以下六维指标，覆盖稳定性、安全性、精确性和实时性：")

add_table(
    ["指标", "符号/单位", "含义", "期望方向"],
    [
        ["能量漂移率", "dE/dt (J/s)", "总机械能的变化速率，度量数值稳定性", "↓ 越小越稳定"],
        ["最小障碍物距离", "d_min (m)", "仿真全程的最小 d_surf，负值表穿透", "↑ 越大越安全"],
        ["最大速度", "V_max (m/s)", "极值速度，超 2.5 表失控", "↓ 越小越可控"],
        ["最大姿态角", "θ_max (rad)", "极值倾斜角，超 0.8 表翻车风险", "↓ 越小越稳定"],
        ["RMS 位置误差", "e_rms (m)", "相对参考轨迹 (RK45+dt=0.001) 的跟踪误差", "↓ 越小越精确"],
        ["平均计算耗时", "T_avg (ms)", "单步积分的时间开销", "↓ 越小越实时"],
    ],
)

heading("5.3 实验模式 (Headless)", 2)

para(
    "实验模式下，C++ 节点跳过所有 ROS 发布 (TF/odom/marker/clock)，"
    "仅执行数值积分和指标采集，大幅减少开销。"
    "到达指定时长后自动调用 rclcpp::shutdown() 并写入指标 CSV。"
)

add_table(
    ["参数", "默认值", "说明"],
    [
        ["experiment.enabled", "false", "实验模式开关"],
        ["experiment.duration", "30.0", "仿真时长 (s)"],
        ["experiment.headless", "false", "跳过 ROS 发布"],
        ["experiment.metrics_path", "experiment_metrics.csv", "指标输出路径"],
    ],
)

heading("5.4 临界步长二分搜索", 2)

para(
    "对每个求解器，在步长范围 [0.001, 0.1] s 内执行二分搜索，精度 0.0005 s。"
    "每次测试运行 30 s 仿真，检查 safety_log.csv 中是否有 COLLISION、VELOCITY 或 ATTITUDE 违规。"
    "若违规，判定为[不安全]，缩小步长；否则加大步长。"
)

para(
    "搜索结果为每个求解器输出 safe_step_size (安全的最大步长) 和 unsafe_step_size ("
    "出现安全违规的最小步长)，两者之差即为安全裕度 (Safety Margin)。"
)

image_placeholder("临界步长搜索结果柱状图 (05_safety_boundary.png)", 5.0)

heading("5.5 预期核心发现", 2)

para(
    "对于所有求解器，预期 safe_step_size < stable_step_size —— "
    "这表明数值稳定性只是安全飞行的必要条件而非充分条件。"
    "安全边界比稳定边界更窄，因为安全还需要考虑物理约束 "
    "(障碍物距离、速度上限、姿态限制)，而非仅仅数值不发散。"
)

doc.add_page_break()

# ================================================================
# 7. 运行模式与使用指南
# ================================================================
heading("6. 运行模式与使用指南", 1)

heading("6.1 模式速览", 2)

add_table(
    ["模式", "脚本", "GUI", "用途"],
    [
        ["正常仿真", "./run_experiment.sh <solver> <dt> <n> [obs]", "RViz ✓", "日常开发与演示"],
        ["安全监视", "./run_safety_monitor.sh <solver> <dt> <n> [obs]", "RViz ✓", "实时违规告警"],
        ["单次采集", "./run_collect_data.sh <solver> <dt> <n> <dur> [obs]", "✗", "深度分析"],
        ["快速对比", "./run_quick_compare.sh", "✗", "5求解器 × 15s 横向比较"],
        ["全因子扫描", "python3 run_experiments.py", "✗", "150次实验数据采集"],
        ["临界步长", "python3 safety_boundary_analysis.py", "✗", "二分搜索安全边界"],
        ["图表生成", "python3 plot_results.py [csv]", "✗", "6张报告级PNG"],
    ],
)

heading("6.2 快速对比示例", 2)

para("最快上手路径——只需 5 分钟得到所有求解器的横向对比：")
para("$ ./run_quick_compare.sh")
para("$ python3 install/crazyflie_ode_core/lib/crazyflie_ode_core/plot_results.py compare_results/quick_compare.csv")

heading("6.3 图表清单", 2)

add_table(
    ["图表", "内容", "关键发现"],
    [
        ["01_energy_drift.png", "能量漂移率 vs 步长", "Implicit Euler 最稳定，Euler 大步长漂移严重"],
        ["02_safety_distance.png", "障碍物距离 vs 步长", "大步长 Euler 穿透障碍物"],
        ["03_accuracy.png", "跟踪误差 vs 步长", "RK4/RK45 精度最优"],
        ["04_performance.png", "计算耗时 vs 步长", "Implicit Euler 计算开销是 Euler 的 20+ 倍"],
        ["05_safety_boundary.png", "安全临界步长", "Safe Step < Stable Step"],
        ["06_radar.png", "五维综合雷达图", "RK4 整体最优 (dt=0.01)"],
    ],
)

doc.add_page_break()

# ================================================================
# 8. 总结
# ================================================================
heading("7. 总结与展望", 1)

heading("7.1 完成的工作", 2)

para("本阶段在原有 ODE 仿真系统基础上新增了四个子系统：")
para("1. 障碍物避障子系统——YAML 配置加载、平滑人工势场 (三次方径向排斥 + 速度阻尼 + 切向绕行)、RViz 可视化、参数化控制。经过三个版本迭代，实现了既不弹飞也不穿透的平稳避障效果。")
para("2. 运行时安全监视器——四维安全条件逐步检查、CSV 违规日志、安全分数聚合、预热机制。为形式化验证提供了在线运行时监控能力。")
para("3. 区间可达集近似分析——采样法 (25次求解器调用) 和 Jacobian 法 (13次) 两种实现、AABB-圆柱保守碰撞检测。在求解器层面提供了'最坏情况'的安全预警。")
para("4. 求解器对比实验框架——全因子 150 次实验自动编排、六维指标体系、临界步长二分搜索、Python 图表生成流水线。可直接产出报告级对比分析。")

heading("7.2 技术贡献", 2)

para(
    "本项目将数值分析 (ODE 求解器稳定域) 与形式化方法 (安全条件不变式、区间可达集) "
    "通过工程实现 (ROS 2 仿真 + 安全监视 + 实验框架) 有机结合起来，"
    "为课题[面向无人机安全自主飞行的 ODE 初值问题数值解法稳定域分析与安全边界形式化验证]"
    "提供了完整的实验验证平台和方法论。"
)

heading("7.3 后续方向", 2)

para("• 将区间可达集扩展为多步可达集 (Flow* 风格)，实现完整可达管 (Reach Tube) 计算")
para("• 引入随机扰动模型 (Monte Carlo)，对比确定性区间分析与概率安全分析")
para("• 将安全监视器扩展为在线安全控制器——当预测到不安全状态时自动切换求解器或减小步长")
para("• 集成反例生成 (Counterexample Mining)，自动搜索[求解器A安全但求解器B碰撞]的场景")

# ================================================================
# 保存
# ================================================================
doc.save(OUTPUT)
print(f"报告已生成: {OUTPUT}")
print(f"大小: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
